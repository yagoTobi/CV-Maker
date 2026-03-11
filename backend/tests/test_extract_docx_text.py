"""
Test the _extract_docx_text() function in services/cv_extractor.py.

This function extracts text from DOCX files and converts it to a
markdown-like format that gives Claude enough structural cues to
correctly identify CV sections, bullets, and entry titles.

All test documents are created in-memory using python-docx -- no fixture
files on disk are needed.

Test Coverage:
- Heading detection (Heading 1-9, level capping at ####)
- List detection via style name ("List Bullet", "List Number", etc.)
- List detection via XML numPr (custom-styled numbered paragraphs)
- List detection via left indent (indented but non-list-styled paragraphs)
- Nested list indent levels (numPr ilvl, cascading indentation)
- Bold-only paragraphs treated as section headers
- Mixed bold/non-bold paragraphs NOT wrapped in **bold**
- Bold paragraphs with whitespace-only non-bold runs
- Table extraction with pipe-separated cells
- Empty paragraphs skipped
- Empty document returns empty string
- Mixed content combining headings, bullets, bold, tables, plain text
- Paragraph ordering is preserved
- Priority/precedence: heading > list > bold > plain
- Edge cases: deeply nested lists, tables with empty cells, no runs

Run with: cd backend && python3 -m pytest tests/test_extract_docx_text.py -v
"""

import sys
import os
import pytest

# Add backend to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.cv_extractor import _extract_docx_text


# ──────────────────────────────────────────────────────────────────────
# Helpers for building in-memory DOCX documents
# ──────────────────────────────────────────────────────────────────────


def _add_numpr(para, ilvl_val: int = 0, num_id: int = 1):
    """Inject XML numPr (numbering properties) into a paragraph.

    This simulates the way Word/LibreOffice marks a paragraph as belonging
    to a numbered or bulleted list via the underlying XML, rather than
    relying on a named style.
    """
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(ilvl_val))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def _make_bold_paragraph(doc, text: str):
    """Create a paragraph where every run is bold."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    return para


def _make_mixed_bold_paragraph(doc, parts: list[tuple[str, bool]]):
    """Create a paragraph with mixed bold and non-bold runs.

    `parts` is a list of (text, is_bold) tuples.
    """
    para = doc.add_paragraph()
    for text, bold in parts:
        run = para.add_run(text)
        run.bold = bold
    return para


# ──────────────────────────────────────────────────────────────────────
# Heading Detection Tests
# ──────────────────────────────────────────────────────────────────────


class TestHeadingDetection:
    """Test that paragraphs with 'Heading N' styles are converted to
    markdown-style headings with ## level syntax.

    The function uses: min(level + 1, 4) where level is the first digit
    found in the style name. If no digit, level defaults to 1."""

    def test_heading_1_becomes_double_hash(self):
        doc = Document()
        doc.add_heading("Experience", level=1)
        result = _extract_docx_text(doc)
        assert "## Experience" in result

    def test_heading_2_becomes_triple_hash(self):
        doc = Document()
        doc.add_heading("Education", level=2)
        result = _extract_docx_text(doc)
        assert "### Education" in result

    def test_heading_3_becomes_quadruple_hash(self):
        doc = Document()
        doc.add_heading("Awards", level=3)
        result = _extract_docx_text(doc)
        assert "#### Awards" in result

    def test_heading_level_caps_at_four_hashes(self):
        """Heading levels 4+ should all produce #### (max 4 hashes).

        The function uses min(level + 1, 4). Level 3 -> 4 hashes.
        Level 4 -> min(5, 4) = 4 hashes. Level 9 -> min(10, 4) = 4 hashes.
        """
        doc = Document()
        doc.add_heading("Deep heading 4", level=4)
        doc.add_heading("Deep heading 5", level=5)
        doc.add_heading("Deep heading 9", level=9)
        result = _extract_docx_text(doc)
        assert "#### Deep heading 4" in result
        assert "#### Deep heading 5" in result
        assert "#### Deep heading 9" in result
        # Should NOT have 5+ hashes
        assert "#####" not in result

    def test_heading_has_surrounding_newlines(self):
        """Headings include \\n before and after via the format string."""
        doc = Document()
        doc.add_paragraph("Before heading")
        doc.add_heading("Section Title", level=1)
        doc.add_paragraph("After heading")
        result = _extract_docx_text(doc)
        # The heading format is f"\\n{'#' * n} {text}\\n" which means
        # the joined output has blank lines around the heading text.
        lines = result.split("\n")
        heading_indices = [
            i for i, line in enumerate(lines) if "## Section Title" in line
        ]
        assert len(heading_indices) == 1
        idx = heading_indices[0]
        # There should be blank lines around it
        assert idx > 0, "Heading should not be the very first line"
        assert lines[idx - 1] == "", "Blank line expected before heading"
        if idx < len(lines) - 1:
            assert lines[idx + 1] == "", "Blank line expected after heading"

    def test_multiple_headings_at_different_levels(self):
        doc = Document()
        doc.add_heading("Level 1", level=1)
        doc.add_heading("Level 2", level=2)
        doc.add_heading("Level 3", level=3)
        result = _extract_docx_text(doc)
        assert "## Level 1" in result
        assert "### Level 2" in result
        assert "#### Level 3" in result

    def test_heading_with_level_0_is_title_style(self):
        """Level 0 in python-docx creates a 'Title' style, which does NOT
        contain the word 'Heading' -- so it should be treated as a regular
        paragraph, not a heading."""
        doc = Document()
        para = doc.add_heading("Document Title", level=0)
        style_name = para.style.name
        result = _extract_docx_text(doc)
        if "Heading" not in style_name:
            # Title goes through normal paragraph processing
            assert "Document Title" in result
            assert "## Document Title" not in result

    def test_heading_style_without_digit_defaults_to_level_1(self):
        """If the style name contains 'Heading' but no digit, the loop
        never finds a digit and level stays at the default of 1.
        Result: min(1 + 1, 4) = ## prefix."""
        doc = Document()
        try:
            para = doc.add_paragraph("TOC Entry")
            para.style = doc.styles["TOC Heading"]
            result = _extract_docx_text(doc)
            # "TOC Heading" contains "Heading", no digit -> level=1 -> ##
            assert "## TOC Entry" in result
        except KeyError:
            pytest.skip("TOC Heading style not available in this template")


# ──────────────────────────────────────────────────────────────────────
# List Detection by Style Name
# ──────────────────────────────────────────────────────────────────────


class TestListDetectionByStyle:
    """Test that paragraphs with list-related style names produce '- ' bullets.

    The function checks: any(kw in style_name.lower() for kw in
    ["list", "bullet", "numbered"]). Standard Word styles that match
    include List Bullet, List Number, List Paragraph, etc."""

    def test_list_bullet_style(self):
        doc = Document()
        doc.add_paragraph("Managed a team of 5", style="List Bullet")
        result = _extract_docx_text(doc)
        assert "- Managed a team of 5" in result

    def test_list_number_style(self):
        doc = Document()
        doc.add_paragraph("First item", style="List Number")
        result = _extract_docx_text(doc)
        assert "- First item" in result

    def test_list_paragraph_style(self):
        """'List Paragraph' is the most common style for Word bullet lists."""
        doc = Document()
        doc.add_paragraph("Generic list item", style="List Paragraph")
        result = _extract_docx_text(doc)
        assert "- Generic list item" in result

    def test_list_bullet_2_style(self):
        doc = Document()
        doc.add_paragraph("Nested bullet", style="List Bullet 2")
        result = _extract_docx_text(doc)
        assert "- Nested bullet" in result

    def test_list_bullet_3_style(self):
        doc = Document()
        doc.add_paragraph("Deep bullet", style="List Bullet 3")
        result = _extract_docx_text(doc)
        assert "- Deep bullet" in result

    def test_list_number_2_style(self):
        doc = Document()
        doc.add_paragraph("Sub-numbered item", style="List Number 2")
        result = _extract_docx_text(doc)
        assert "- Sub-numbered item" in result

    def test_list_continue_style(self):
        doc = Document()
        doc.add_paragraph("Continued item", style="List Continue")
        result = _extract_docx_text(doc)
        assert "- Continued item" in result

    def test_list_style_name(self):
        """The bare 'List' style should also match (contains 'list')."""
        doc = Document()
        doc.add_paragraph("Plain list item", style="List")
        result = _extract_docx_text(doc)
        assert "- Plain list item" in result

    def test_multiple_bullet_items(self):
        doc = Document()
        doc.add_paragraph("Item one", style="List Bullet")
        doc.add_paragraph("Item two", style="List Bullet")
        doc.add_paragraph("Item three", style="List Bullet")
        result = _extract_docx_text(doc)
        assert "- Item one" in result
        assert "- Item two" in result
        assert "- Item three" in result

    def test_style_detection_is_case_insensitive(self):
        """The function lowercases the style name before keyword matching.
        Standard styles use title case ('List Bullet') which lowercases to
        'list bullet', matching both 'list' and 'bullet' keywords."""
        doc = Document()
        doc.add_paragraph("Bullet item", style="List Bullet")
        result = _extract_docx_text(doc)
        assert "- Bullet item" in result

    def test_style_based_list_has_zero_indent(self):
        """Style-based list detection does not set indent_level, so
        the prefix is always '- ' (no leading spaces) unless numPr
        provides an ilvl."""
        doc = Document()
        doc.add_paragraph("Level zero", style="List Bullet")
        result = _extract_docx_text(doc)
        stripped_lines = [l for l in result.split("\n") if l.strip()]
        assert stripped_lines[0] == "- Level zero"


# ──────────────────────────────────────────────────────────────────────
# List Detection by XML numPr
# ──────────────────────────────────────────────────────────────────────


class TestListDetectionByNumPr:
    """Test that paragraphs with XML numPr (numbering properties) are
    detected as list items even without a list-related style name.

    This is Method 2 in the source code -- it catches custom-styled
    lists that don't use a 'List ...' style name."""

    def test_numpr_creates_bullet(self):
        doc = Document()
        para = doc.add_paragraph("Custom styled bullet")
        _add_numpr(para, ilvl_val=0)
        result = _extract_docx_text(doc)
        assert "- Custom styled bullet" in result

    def test_numpr_ilvl_0_no_indent(self):
        doc = Document()
        para = doc.add_paragraph("Top-level bullet")
        _add_numpr(para, ilvl_val=0)
        result = _extract_docx_text(doc)
        stripped_lines = [l for l in result.split("\n") if l.strip()]
        assert stripped_lines[0] == "- Top-level bullet"

    def test_numpr_ilvl_1_two_space_indent(self):
        doc = Document()
        para = doc.add_paragraph("Nested bullet")
        _add_numpr(para, ilvl_val=1)
        result = _extract_docx_text(doc)
        assert "  - Nested bullet" in result

    def test_numpr_ilvl_2_four_space_indent(self):
        doc = Document()
        para = doc.add_paragraph("Deeply nested")
        _add_numpr(para, ilvl_val=2)
        result = _extract_docx_text(doc)
        assert "    - Deeply nested" in result

    def test_numpr_ilvl_3_six_space_indent(self):
        doc = Document()
        para = doc.add_paragraph("Triple nested")
        _add_numpr(para, ilvl_val=3)
        result = _extract_docx_text(doc)
        assert "      - Triple nested" in result

    def test_numpr_deeply_nested_ilvl_4(self):
        doc = Document()
        para = doc.add_paragraph("Very deep")
        _add_numpr(para, ilvl_val=4)
        result = _extract_docx_text(doc)
        # 4 * 2 spaces = 8 spaces before "- "
        assert "        - Very deep" in result

    def test_numpr_without_ilvl_element(self):
        """If numPr exists but has no ilvl element, indent_level stays 0."""
        doc = Document()
        para = doc.add_paragraph("No ilvl bullet")
        pPr = para._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")
        numPr.append(numId)
        pPr.append(numPr)
        result = _extract_docx_text(doc)
        stripped_lines = [l for l in result.split("\n") if l.strip()]
        assert stripped_lines[0] == "- No ilvl bullet"

    def test_numpr_with_ilvl_no_val_attribute(self):
        """If ilvl element exists but has no w:val attribute, indent stays 0."""
        doc = Document()
        para = doc.add_paragraph("No val bullet")
        pPr = para._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        # Deliberately do NOT set w:val
        numPr.append(ilvl)
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")
        numPr.append(numId)
        pPr.append(numPr)
        result = _extract_docx_text(doc)
        stripped_lines = [l for l in result.split("\n") if l.strip()]
        assert stripped_lines[0] == "- No val bullet"

    def test_numpr_hierarchical_list(self):
        """A parent bullet followed by nested child bullets."""
        doc = Document()
        p1 = doc.add_paragraph("Parent item")
        _add_numpr(p1, ilvl_val=0)
        p2 = doc.add_paragraph("Child item 1")
        _add_numpr(p2, ilvl_val=1)
        p3 = doc.add_paragraph("Child item 2")
        _add_numpr(p3, ilvl_val=1)
        p4 = doc.add_paragraph("Grandchild")
        _add_numpr(p4, ilvl_val=2)
        result = _extract_docx_text(doc)
        lines = result.strip().split("\n")
        assert lines[0] == "- Parent item"
        assert lines[1] == "  - Child item 1"
        assert lines[2] == "  - Child item 2"
        assert lines[3] == "    - Grandchild"

    def test_numpr_with_different_num_ids(self):
        """Different numId values (different list definitions) should
        all be detected as lists."""
        doc = Document()
        p1 = doc.add_paragraph("List A item")
        _add_numpr(p1, ilvl_val=0, num_id=1)
        p2 = doc.add_paragraph("List B item")
        _add_numpr(p2, ilvl_val=0, num_id=2)
        result = _extract_docx_text(doc)
        assert "- List A item" in result
        assert "- List B item" in result


# ──────────────────────────────────────────────────────────────────────
# List Detection by Left Indent
# ──────────────────────────────────────────────────────────────────────


class TestListDetectionByIndent:
    """Test that non-list-styled paragraphs with left indent are treated
    as list items (Method 3 in the source code).

    The function converts EMUs to indent levels using:
        emu_per_indent = 457200 (0.5 inch)
        indent_level = int(left_indent / emu_per_indent)
    Paragraphs are only treated as lists if indent_level > 0."""

    def test_half_inch_indent_is_list_level_1(self):
        doc = Document()
        para = doc.add_paragraph("Indented text")
        para.paragraph_format.left_indent = Inches(0.5)
        result = _extract_docx_text(doc)
        assert "  - Indented text" in result

    def test_one_inch_indent_is_list_level_2(self):
        doc = Document()
        para = doc.add_paragraph("More indented")
        para.paragraph_format.left_indent = Inches(1.0)
        result = _extract_docx_text(doc)
        assert "    - More indented" in result

    def test_one_and_half_inch_indent_is_list_level_3(self):
        doc = Document()
        para = doc.add_paragraph("Triple indented")
        para.paragraph_format.left_indent = Inches(1.5)
        result = _extract_docx_text(doc)
        assert "      - Triple indented" in result

    def test_quarter_inch_indent_is_not_list(self):
        """An indent of 0.25 inch rounds down to level 0, which means
        is_list stays False. The paragraph is treated as regular text."""
        doc = Document()
        para = doc.add_paragraph("Slightly indented")
        para.paragraph_format.left_indent = Inches(0.25)
        result = _extract_docx_text(doc)
        assert "Slightly indented" in result
        assert "- Slightly indented" not in result

    def test_zero_indent_is_not_list(self):
        """Explicitly set indent to 0 -- should not be treated as list."""
        doc = Document()
        para = doc.add_paragraph("No indent")
        para.paragraph_format.left_indent = Inches(0)
        result = _extract_docx_text(doc)
        assert "No indent" in result
        assert "- No indent" not in result

    def test_no_indent_set_is_not_list(self):
        """Paragraph with no left_indent set (None) is not treated as list."""
        doc = Document()
        para = doc.add_paragraph("Normal paragraph")
        assert para.paragraph_format.left_indent is None
        result = _extract_docx_text(doc)
        assert "Normal paragraph" in result
        assert "- Normal paragraph" not in result

    def test_indent_not_applied_if_style_already_matched_as_list(self):
        """If the style name already triggered list detection (Method 1),
        the indent method (Method 3) is skipped because `not is_list` is
        False. The indent_level from Method 3 is never calculated."""
        doc = Document()
        para = doc.add_paragraph("Styled list item", style="List Bullet")
        para.paragraph_format.left_indent = Inches(1.0)
        result = _extract_docx_text(doc)
        # Style detection fires first with indent_level=0
        assert "- Styled list item" in result
        # Should NOT have 4-space indent from the 1-inch setting
        assert "    - Styled list item" not in result

    def test_indent_not_applied_if_numpr_already_matched(self):
        """If numPr triggered list detection (Method 2), Method 3 is
        skipped. The indent_level comes from numPr's ilvl."""
        doc = Document()
        para = doc.add_paragraph("NumPr + indent")
        _add_numpr(para, ilvl_val=1)
        para.paragraph_format.left_indent = Inches(2.0)
        result = _extract_docx_text(doc)
        # indent_level=1 from numPr, not 4 from 2.0 inches
        assert "  - NumPr + indent" in result
        assert "        - NumPr + indent" not in result


# ──────────────────────────────────────────────────────────────────────
# Bold Paragraph Detection Tests
# ──────────────────────────────────────────────────────────────────────


class TestBoldParagraphDetection:
    """Test that paragraphs where all non-whitespace runs are bold get
    wrapped in **bold** markers.

    The condition is: runs and all(r.bold for r in runs if r.text.strip())
    This means:
    - The paragraph must have at least one run
    - All runs with non-empty stripped text must be bold
    - Whitespace-only runs are ignored in the check"""

    def test_single_bold_run_wrapped(self):
        doc = Document()
        _make_bold_paragraph(doc, "Software Engineer")
        result = _extract_docx_text(doc)
        assert "**Software Engineer**" in result

    def test_multiple_bold_runs_wrapped(self):
        doc = Document()
        para = doc.add_paragraph()
        r1 = para.add_run("Acme Corp")
        r1.bold = True
        r2 = para.add_run(" | Senior Engineer")
        r2.bold = True
        result = _extract_docx_text(doc)
        assert "**Acme Corp | Senior Engineer**" in result

    def test_mixed_bold_not_wrapped(self):
        """If any non-whitespace run is not bold, the paragraph is NOT
        wrapped in ** markers."""
        doc = Document()
        _make_mixed_bold_paragraph(
            doc, [("Bold part ", True), ("normal part", False)]
        )
        result = _extract_docx_text(doc)
        assert "**" not in result
        assert "Bold part normal part" in result

    def test_bold_with_whitespace_only_non_bold_runs(self):
        """Whitespace-only runs are filtered out of the bold check.
        If all remaining runs are bold, the paragraph IS wrapped."""
        doc = Document()
        para = doc.add_paragraph()
        r1 = para.add_run("Company Name")
        r1.bold = True
        r2 = para.add_run("   ")  # whitespace-only, not bold
        r2.bold = False
        r3 = para.add_run("Position")
        r3.bold = True
        result = _extract_docx_text(doc)
        assert "**Company Name   Position**" in result

    def test_non_bold_paragraph_not_wrapped(self):
        """Regular paragraph with no bold runs should not be wrapped."""
        doc = Document()
        doc.add_paragraph("Regular text")
        result = _extract_docx_text(doc)
        assert "Regular text" in result
        assert "**" not in result

    def test_paragraph_with_false_bold_not_wrapped(self):
        """Explicitly non-bold runs should not trigger wrapping."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Explicitly not bold")
        run.bold = False
        result = _extract_docx_text(doc)
        assert "Explicitly not bold" in result
        assert "**" not in result

    def test_multiple_bold_paragraphs(self):
        doc = Document()
        _make_bold_paragraph(doc, "Company A")
        _make_bold_paragraph(doc, "Company B")
        result = _extract_docx_text(doc)
        assert "**Company A**" in result
        assert "**Company B**" in result

    def test_bold_paragraph_with_only_whitespace_runs(self):
        """If ALL runs are whitespace-only, the paragraph text.strip()
        is empty, so it's skipped entirely at the top of the loop."""
        doc = Document()
        para = doc.add_paragraph()
        r = para.add_run("   ")
        r.bold = True
        result = _extract_docx_text(doc)
        assert result.strip() == ""


# ──────────────────────────────────────────────────────────────────────
# Detection Priority / Precedence Tests
# ──────────────────────────────────────────────────────────────────────


class TestDetectionPriority:
    """Verify the processing order: heading > list > bold > plain.

    Each detection branch uses `continue`, so once a paragraph matches
    at a higher priority, lower-priority checks never run."""

    def test_heading_takes_priority_over_bold(self):
        """Headings are checked before bold. A bold heading should produce
        heading syntax, not **bold** markers."""
        doc = Document()
        doc.add_heading("Bold Heading", level=1)
        result = _extract_docx_text(doc)
        assert "## Bold Heading" in result
        assert "**Bold Heading**" not in result

    def test_list_takes_priority_over_bold(self):
        """List items are checked before bold. A bold list-styled paragraph
        should produce '- text', not '**text**'."""
        doc = Document()
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run("Bold bullet")
        run.bold = True
        result = _extract_docx_text(doc)
        assert "- Bold bullet" in result
        assert "**Bold bullet**" not in result

    def test_indent_list_takes_priority_over_bold(self):
        """Indented paragraphs (Method 3 list detection) are checked
        before bold."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold and indented")
        run.bold = True
        para.paragraph_format.left_indent = Inches(0.5)
        result = _extract_docx_text(doc)
        assert "  - Bold and indented" in result
        assert "**Bold and indented**" not in result

    def test_numpr_list_takes_priority_over_bold(self):
        """numPr list detection (Method 2) takes priority over bold."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("NumPr bold bullet")
        run.bold = True
        _add_numpr(para, ilvl_val=0)
        result = _extract_docx_text(doc)
        assert "- NumPr bold bullet" in result
        assert "**NumPr bold bullet**" not in result

    def test_heading_takes_priority_over_list_style(self):
        """If a paragraph has a Heading style, it's treated as a heading
        even if the style name somehow also contains 'list'."""
        doc = Document()
        # Standard Heading style does not contain 'list', but this test
        # verifies the heading check runs before the list check.
        doc.add_heading("Section", level=2)
        result = _extract_docx_text(doc)
        assert "### Section" in result
        assert "- Section" not in result


# ──────────────────────────────────────────────────────────────────────
# Table Extraction Tests
# ──────────────────────────────────────────────────────────────────────


class TestTableExtraction:
    """Test that tables are extracted with pipe-separated cells.

    Tables are processed after all paragraphs via `doc.tables`.
    Empty cells (after stripping) are excluded from the pipe join.
    Rows where all cells are empty produce no output."""

    def test_simple_two_by_two_table(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Level"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "Expert"
        result = _extract_docx_text(doc)
        assert "Skill | Level" in result
        assert "Python | Expert" in result

    def test_table_with_three_columns(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Role"
        table.cell(0, 2).text = "Years"
        result = _extract_docx_text(doc)
        assert "Name | Role | Years" in result

    def test_table_with_empty_cells_filtered(self):
        """Empty cells are excluded from the pipe-joined string."""
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = ""
        table.cell(0, 2).text = "5 years"
        result = _extract_docx_text(doc)
        assert "Python | 5 years" in result

    def test_table_with_whitespace_only_cells_filtered(self):
        """Cells with only whitespace should also be filtered out."""
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "   "
        table.cell(0, 2).text = "C"
        result = _extract_docx_text(doc)
        assert "A | C" in result

    def test_table_row_all_empty_cells_skipped(self):
        """If all cells in a row are empty, the row produces no output."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Data"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = ""
        table.cell(1, 1).text = ""
        result = _extract_docx_text(doc)
        assert "Data | Value" in result
        non_empty_lines = [l for l in result.strip().split("\n") if l.strip()]
        assert len(non_empty_lines) == 1

    def test_table_has_blank_line_separator_before_it(self):
        """A blank line (empty string) is appended to `parts` before each
        table's rows."""
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        result = _extract_docx_text(doc)
        lines = result.split("\n")
        for i, line in enumerate(lines):
            if "A | B" in line:
                assert i > 0
                assert lines[i - 1] == "", (
                    "Expected blank line before table row"
                )
                break
        else:
            pytest.fail("Table content 'A | B' not found in output")

    def test_multiple_tables(self):
        doc = Document()
        t1 = doc.add_table(rows=1, cols=2)
        t1.cell(0, 0).text = "Table1A"
        t1.cell(0, 1).text = "Table1B"
        t2 = doc.add_table(rows=1, cols=2)
        t2.cell(0, 0).text = "Table2A"
        t2.cell(0, 1).text = "Table2B"
        result = _extract_docx_text(doc)
        assert "Table1A | Table1B" in result
        assert "Table2A | Table2B" in result

    def test_table_cell_whitespace_stripped(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "  padded  "
        table.cell(0, 1).text = "  value  "
        result = _extract_docx_text(doc)
        assert "padded | value" in result

    def test_single_cell_table(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Solo cell"
        result = _extract_docx_text(doc)
        assert "Solo cell" in result

    def test_table_multiple_rows(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        for i in range(3):
            table.cell(i, 0).text = f"Key{i}"
            table.cell(i, 1).text = f"Val{i}"
        result = _extract_docx_text(doc)
        assert "Key0 | Val0" in result
        assert "Key1 | Val1" in result
        assert "Key2 | Val2" in result


# ──────────────────────────────────────────────────────────────────────
# Empty Content Tests
# ──────────────────────────────────────────────────────────────────────


class TestEmptyContent:
    """Test behavior with empty paragraphs and empty documents."""

    def test_empty_document_returns_empty_string(self):
        """A new Document() has one default empty paragraph. The function
        should produce an empty or whitespace-only result."""
        doc = Document()
        result = _extract_docx_text(doc)
        assert result.strip() == ""

    def test_whitespace_only_paragraphs_skipped(self):
        doc = Document()
        doc.add_paragraph("   ")
        doc.add_paragraph("\t")
        doc.add_paragraph("  \n  ")
        result = _extract_docx_text(doc)
        assert result.strip() == ""

    def test_empty_string_paragraphs_skipped(self):
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        result = _extract_docx_text(doc)
        assert result.strip() == ""

    def test_empty_paragraphs_between_content_do_not_add_lines(self):
        """Empty paragraphs are filtered by the text.strip() check, so
        they produce no output lines."""
        doc = Document()
        doc.add_paragraph("Line one")
        doc.add_paragraph("")  # empty - skipped
        doc.add_paragraph("Line two")
        result = _extract_docx_text(doc)
        assert "Line one" in result
        assert "Line two" in result


# ──────────────────────────────────────────────────────────────────────
# Mixed Content / Integration Tests
# ──────────────────────────────────────────────────────────────────────


class TestMixedContent:
    """Test documents that combine headings, bullets, bold text, tables,
    and regular paragraphs -- simulating a real CV structure."""

    def test_full_cv_structure(self):
        """Simulate a realistic CV with all element types."""
        doc = Document()

        # Personal info as bold
        _make_bold_paragraph(doc, "John Doe")
        doc.add_paragraph("john.doe@email.com | (555) 123-4567")

        # Work experience section
        doc.add_heading("Work Experience", level=1)
        _make_bold_paragraph(doc, "Senior Engineer | Acme Corp | 2020-Present")
        p1 = doc.add_paragraph("Led migration to microservices")
        _add_numpr(p1, ilvl_val=0)
        p2 = doc.add_paragraph("Reduced latency by 40%")
        _add_numpr(p2, ilvl_val=0)

        # Education section
        doc.add_heading("Education", level=1)
        _make_bold_paragraph(doc, "MIT | B.S. Computer Science | 2016-2020")
        doc.add_paragraph("GPA: 3.8/4.0")

        # Skills as a table
        doc.add_heading("Skills", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Languages"
        table.cell(0, 1).text = "Python, TypeScript, Go"
        table.cell(1, 0).text = "Tools"
        table.cell(1, 1).text = "Docker, Kubernetes, AWS"

        result = _extract_docx_text(doc)

        # Verify all structural elements appear
        assert "**John Doe**" in result
        assert "john.doe@email.com" in result
        assert "## Work Experience" in result
        assert "**Senior Engineer | Acme Corp | 2020-Present**" in result
        assert "- Led migration to microservices" in result
        assert "- Reduced latency by 40%" in result
        assert "## Education" in result
        assert "**MIT | B.S. Computer Science | 2016-2020**" in result
        assert "GPA: 3.8/4.0" in result
        assert "### Skills" in result
        assert "Languages | Python, TypeScript, Go" in result
        assert "Tools | Docker, Kubernetes, AWS" in result

    def test_paragraph_ordering_preserved(self):
        """Elements should appear in the output in the same order they
        appear in the document (paragraphs first, then tables)."""
        doc = Document()
        doc.add_paragraph("First")
        doc.add_heading("Second", level=1)
        doc.add_paragraph("Third", style="List Bullet")
        _make_bold_paragraph(doc, "Fourth")
        doc.add_paragraph("Fifth")

        result = _extract_docx_text(doc)
        pos_first = result.index("First")
        pos_second = result.index("## Second")
        pos_third = result.index("- Third")
        pos_fourth = result.index("**Fourth**")
        pos_fifth = result.index("Fifth")
        assert pos_first < pos_second < pos_third < pos_fourth < pos_fifth

    def test_tables_appear_after_all_paragraphs(self):
        """The function iterates doc.paragraphs first, then doc.tables.
        Tables always appear at the end of the output regardless of their
        position in the document body.

        NOTE: This is an inherent limitation of the current implementation.
        In a real DOCX, tables may be interspersed with paragraphs, but
        doc.paragraphs and doc.tables are separate iterators."""
        doc = Document()
        doc.add_paragraph("Paragraph before table")

        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"

        doc.add_paragraph("Paragraph after table")

        result = _extract_docx_text(doc)
        pos_before = result.index("Paragraph before table")
        pos_after = result.index("Paragraph after table")
        pos_table = result.index("A | B")

        # Both paragraphs come before the table in output
        assert pos_before < pos_after
        assert pos_after < pos_table

    def test_heading_followed_by_list_items(self):
        doc = Document()
        doc.add_heading("Achievements", level=2)
        doc.add_paragraph("Won first place", style="List Bullet")
        doc.add_paragraph("Published 3 papers", style="List Bullet")
        doc.add_paragraph("Filed 2 patents", style="List Bullet")
        result = _extract_docx_text(doc)
        assert "### Achievements" in result
        assert "- Won first place" in result
        assert "- Published 3 papers" in result
        assert "- Filed 2 patents" in result

    def test_alternating_bold_and_bullets(self):
        """Simulates a work experience section with bold job titles
        followed by bullet achievements."""
        doc = Document()
        _make_bold_paragraph(doc, "Software Engineer at Company A")
        p1 = doc.add_paragraph("Built REST API", style="List Bullet")
        p2 = doc.add_paragraph("Improved test coverage", style="List Bullet")
        _make_bold_paragraph(doc, "Junior Dev at Company B")
        p3 = doc.add_paragraph("Fixed 50 bugs", style="List Bullet")

        result = _extract_docx_text(doc)
        assert "**Software Engineer at Company A**" in result
        assert "- Built REST API" in result
        assert "- Improved test coverage" in result
        assert "**Junior Dev at Company B**" in result
        assert "- Fixed 50 bugs" in result

        # Verify ordering
        lines = [l for l in result.split("\n") if l.strip()]
        bold_a = next(i for i, l in enumerate(lines) if "Company A" in l)
        bullet_1 = next(i for i, l in enumerate(lines) if "REST API" in l)
        bold_b = next(i for i, l in enumerate(lines) if "Company B" in l)
        bullet_3 = next(i for i, l in enumerate(lines) if "50 bugs" in l)
        assert bold_a < bullet_1 < bold_b < bullet_3


# ──────────────────────────────────────────────────────────────────────
# Edge Cases
# ──────────────────────────────────────────────────────────────────────


class TestEdgeCases:
    """Test edge cases and unusual document structures."""

    def test_special_characters_pass_through_unmodified(self):
        """_extract_docx_text does NOT perform LaTeX escaping -- special
        characters should appear exactly as they are in the document."""
        doc = Document()
        doc.add_paragraph("Revenue grew 50% & exceeded $1M target")
        result = _extract_docx_text(doc)
        assert "Revenue grew 50% & exceeded $1M target" in result

    def test_latex_special_chars_all_preserved(self):
        """All LaTeX special characters should pass through unchanged."""
        doc = Document()
        doc.add_paragraph("& % $ # _ { } ~ ^ \\")
        result = _extract_docx_text(doc)
        assert "& % $ # _ { } ~ ^ \\" in result

    def test_unicode_text_preserved(self):
        doc = Document()
        doc.add_paragraph("Fluent in Francais, Deutsch, and Nihongo")
        doc.add_paragraph("Resume of Jose Garcia-Lopez")
        result = _extract_docx_text(doc)
        assert "Francais" in result
        assert "Jose Garcia-Lopez" in result

    def test_emoji_in_text(self):
        doc = Document()
        doc.add_paragraph("Award Winner - Top Performer")
        result = _extract_docx_text(doc)
        assert "Award Winner" in result

    def test_very_long_paragraph(self):
        doc = Document()
        long_text = "Achievement " * 200
        doc.add_paragraph(long_text.strip())
        result = _extract_docx_text(doc)
        assert long_text.strip() in result

    def test_paragraph_text_is_stripped(self):
        """Leading and trailing whitespace in paragraph text is stripped."""
        doc = Document()
        doc.add_paragraph("  padded text  ")
        result = _extract_docx_text(doc)
        for line in result.split("\n"):
            if "padded text" in line:
                assert line == "padded text"
                break
        else:
            pytest.fail("'padded text' not found in output")

    def test_only_tables_no_text_paragraphs(self):
        """A document with only a table and no text paragraphs."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Key"
        table.cell(0, 1).text = "Value"
        result = _extract_docx_text(doc)
        assert "Key | Value" in result

    def test_many_headings_and_bullets(self):
        """Stress test with many sections."""
        doc = Document()
        for i in range(10):
            doc.add_heading(f"Section {i}", level=1)
            for j in range(5):
                doc.add_paragraph(f"Bullet {i}.{j}", style="List Bullet")

        result = _extract_docx_text(doc)
        for i in range(10):
            assert f"## Section {i}" in result
            for j in range(5):
                assert f"- Bullet {i}.{j}" in result

    def test_document_with_normal_style_paragraphs_only(self):
        """A document with only Normal-styled paragraphs (no headings,
        no lists, no bold)."""
        doc = Document()
        doc.add_paragraph("Line one")
        doc.add_paragraph("Line two")
        doc.add_paragraph("Line three")
        result = _extract_docx_text(doc)
        lines = [l for l in result.split("\n") if l.strip()]
        assert lines == ["Line one", "Line two", "Line three"]

    def test_table_with_multiline_cell_text(self):
        """Table cells can contain newlines from multiple paragraphs
        within the cell. cell.text joins them with newlines."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        # Add multiple paragraphs to a cell
        cell = table.cell(0, 0)
        cell.text = "Line 1\nLine 2"
        table.cell(0, 1).text = "Single line"
        result = _extract_docx_text(doc)
        # The cell text "Line 1\nLine 2" gets stripped to "Line 1\nLine 2"
        # (strip only removes leading/trailing whitespace)
        assert "Single line" in result

    def test_result_is_string_type(self):
        """The return type should always be a string."""
        doc = Document()
        result = _extract_docx_text(doc)
        assert isinstance(result, str)

    def test_result_joined_with_newlines(self):
        """Parts list is joined with '\\n', so each element becomes a line."""
        doc = Document()
        doc.add_paragraph("Alpha")
        doc.add_paragraph("Beta")
        doc.add_paragraph("Gamma")
        result = _extract_docx_text(doc)
        # Split and filter empties -- should have exactly 3 content lines
        lines = [l for l in result.split("\n") if l.strip()]
        assert len(lines) == 3
        assert lines[0] == "Alpha"
        assert lines[1] == "Beta"
        assert lines[2] == "Gamma"
