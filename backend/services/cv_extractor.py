"""
CV extraction service for importing CVs from PDF, DOCX, and JSON files.
Uses Claude AI via AWS Bedrock for intelligent extraction from unstructured documents.
"""

import json
import logging
from typing import Optional, List, Dict, Any
from pydantic import BaseModel, ValidationError
from docx import Document
from io import BytesIO

from routes.cv_versions import CVFormData
from services.bedrock import bedrock_client

logger = logging.getLogger(__name__)


class ImportSummary(BaseModel):
    workEntries: int = 0
    educationEntries: int = 0
    skillCategories: int = 0
    projects: int = 0
    awards: int = 0
    additionalSections: int = 0


class CVImportResult(BaseModel):
    success: bool
    form_data: Optional[Dict[str, Any]] = None
    source: str  # "pdf" | "docx" | "json"
    confidence: Optional[Dict[str, Any]] = None
    summary: Optional[ImportSummary] = None
    warnings: Optional[List[str]] = None
    error: Optional[str] = None


EXTRACTION_SYSTEM_PROMPT = """You are a CV/resume data extraction expert. Your task is to extract structured information from CV documents and return it as JSON.

CRITICAL RULES:
1. Return ONLY valid JSON — no markdown fences, no explanations, just the JSON object.
2. Leave fields as empty strings ("") rather than guessing or fabricating information.
3. Normalize all dates to "MMM YYYY" format (e.g., "Jan 2021", "Dec 2023").
4. For current/ongoing positions, use "Present" as the endDate.
5. Extract bullet points / details as arrays of strings. For work experience, education, projects, and additional sections: preserve the original bullet granularity from the source document — each bullet point in the source should map to exactly one string in the bullets array. For the skills section specifically: parse each skill or closely related skill group as a separate array item (see Skills Parsing Instructions below).
6. Preserve the order of entries within each section exactly as they appear in the source document. Work entries, education entries, skills categories, etc. should appear in the same sequence as the original.
7. Include a _confidence annotation block and a _warnings array.

Required JSON schema:

{
  "sectionOrder": ["work", "education", "skills", "projects", "awards", "additional-0"],
  "personalInfo": {
    "fullName": "",
    "email": "",
    "phone": "",
    "location": "",
    "links": [{"label": "string (e.g. LinkedIn, GitHub, Portfolio)", "url": "string"}],
    "summary": ""
  },
  "workExperience": [
    {
      "company": "",
      "title": "",
      "startDate": "MMM YYYY",
      "endDate": "MMM YYYY or Present",
      "location": "",
      "bullets": [""]
    }
  ],
  "education": [
    {
      "school": "",
      "degree": "",
      "startDate": "MMM YYYY",
      "endDate": "MMM YYYY",
      "location": "",
      "gpa": "",
      "details": [""]
    }
  ],
  "skills": [
    {
      "category": "e.g. Programming Languages, Tools, Frameworks",
      "skills": ["individual skill 1", "individual skill 2"]
    }
  ],
  "projects": [
    {
      "name": "",
      "year": "YYYY",
      "description": "",
      "technologies": "",
      "bullets": ["detailed achievement or responsibility"]
    }
  ],
  "awards": [
    {
      "year": "YYYY",
      "title": "",
      "description": ""
    }
  ],
  "additionalSections": [
    {
      "title": "Section title from the document (e.g., Leadership, Certifications, Volunteer Work)",
      "entries": [
        {
          "title": "",
          "subtitle": "",
          "startDate": "MMM YYYY",
          "endDate": "MMM YYYY or Present",
          "location": "",
          "description": "",
          "bullets": [""]
        }
      ]
    }
  ],
  "_confidence": {
    "overall": "high|medium|low",
    "fields": {
      "personalInfo.email": "medium",
      "workExperience[0].endDate": "low"
    }
  },
  "_warnings": []
}

Set sectionOrder to reflect the order that sections appear in the source document. Use these exact keys for standard sections: work, education, skills, projects, awards. For additional sections, use "additional-0", "additional-1", etc., corresponding to their index in the additionalSections array. Only include sections that have data.

Map standard CV sections to their dedicated fields (workExperience, education, skills, projects, awards). For any section that does NOT fit these standard types — such as Leadership, Extra Curricular Activities, Certifications, Volunteer Work, Publications, Research, Languages, Hobbies, or any other custom section — place it in additionalSections with the original section title preserved. NEVER silently drop content that doesn't fit the standard sections.

If a project has detailed bullet points or achievements, include them in the bullets array. Use the description field for a brief summary and bullets for detailed points. If there's only a description with no bullets, leave bullets as an empty array.

Skills Parsing Instructions:
For the skills array, intelligently parse and group skills into logical categories:
- If the CV lists skills as bullet points, extract EACH skill as a separate string in the skills array.
- If a bullet contains multiple related items separated by commas or "and", split them into individual skills.
- Common skill categories: "Programming Languages", "Frameworks & Libraries", "Tools & Technologies", "Databases", "Cloud Platforms", "Soft Skills", "Languages", "Certifications".
- Remove filler words like "Knowledge of", "Experience with", "Proven skills in" — just keep the skill itself.
- For technical acronyms grouped together (e.g., "TCP/IP, DNS, DHCP"), each protocol or technology should be a separate item unless they are a tightly coupled pair (e.g., "TCP/IP" stays as one item).
- Examples:
  - "Knowledge of Python, Java, and C++" -> ["Python", "Java", "C++"]
  - "Experience with React and Vue frameworks" -> ["React", "Vue"]
  - "Excellent communication and presentation skills" -> ["Communication skills", "Presentation skills"]
  - "AWS, Azure, and GCP" -> ["AWS", "Azure", "GCP"]
  - "Routing protocols (BGP/OSPF, MPLS)" -> ["BGP", "OSPF", "MPLS"] or ["Routing protocols (BGP/OSPF)", "MPLS"] if context groups them
  - "Good organisational skills with the ability to work as part of a team" -> ["Organisational skills", "Teamwork"]

Confidence guidelines:
- "high": clearly stated and unambiguous
- "medium": required some interpretation or inference
- "low": unclear, incomplete, or uncertain
In the _confidence.fields map, only include fields with "medium" or "low" confidence. Omit fields that are "high" confidence — high is the assumed default. This saves output space for actual CV content.

Date examples:
- "January 2021" → "Jan 2021"
- "01/2021" → "Jan 2021"
- "2021-01" → "Jan 2021"
- "Current" / "Now" / "present" → "Present"

For links, identify: LinkedIn, GitHub, personal websites, portfolios.
Omit sections entirely (empty array) if the CV has no data for them."""

EXTRACTION_USER_PROMPT = (
    "Extract all CV/resume information from this document and return it as "
    "structured JSON matching the schema in your instructions. "
    "Be thorough but accurate — only include information actually present."
)


def _parse_extraction_response(raw: str, source: str) -> CVImportResult:
    """Parse Claude's JSON response and build a CVImportResult."""
    # Strip markdown fences if Claude wrapped the response
    text = raw.strip()
    if text.startswith("```"):
        first_newline = text.index("\n")
        text = text[first_newline + 1:]
    if text.endswith("```"):
        text = text[:-3].rstrip()

    # Detect truncation — response likely hit max_tokens
    truncated = not text.rstrip().endswith("}")

    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        if truncated:
            logger.warning("Extraction response appears truncated (did not end with })")
            return CVImportResult(
                success=False,
                source=source,
                error="The extraction was cut short — your CV may be too long for a single pass. Try removing some content or uploading a shorter version.",
            )
        logger.error("Failed to parse extraction JSON: %s — raw[:500]: %s", e, raw[:500])
        return CVImportResult(
            success=False,
            source=source,
            error="Failed to parse extraction results. The AI response was not valid JSON.",
        )

    confidence = data.pop("_confidence", {"overall": "medium", "fields": {}})
    warnings: list[str] = data.pop("_warnings", [])
    data.pop("templateId", None)

    summary = ImportSummary(
        workEntries=len(data.get("workExperience", [])),
        educationEntries=len(data.get("education", [])),
        skillCategories=len(data.get("skills", [])),
        projects=len(data.get("projects") or []),
        awards=len(data.get("awards") or []),
        additionalSections=len(data.get("additionalSections") or []),
    )

    # Validate loosely — try building a CVFormData to catch shape issues
    try:
        CVFormData(**(data | {"templateId": "_validation"}))
    except ValidationError as e:
        logger.warning("Extracted data has schema issues: %s", e)
        warnings.append("Some fields may not match the expected format.")

    return CVImportResult(
        success=True,
        form_data=data,
        source=source,
        confidence=confidence,
        summary=summary,
        warnings=warnings or None,
    )


async def extract_from_pdf(file_bytes: bytes) -> CVImportResult:
    """Send PDF to Claude as a document content block for multimodal extraction."""
    try:
        response = bedrock_client.chat_with_document(
            document_bytes=file_bytes,
            document_media_type="application/pdf",
            text_prompt=EXTRACTION_USER_PROMPT,
            system_prompt=EXTRACTION_SYSTEM_PROMPT,
            max_tokens=8192,
        )
        return _parse_extraction_response(response, source="pdf")
    except Exception as e:
        logger.error("PDF extraction failed: %s", e)
        return CVImportResult(
            success=False,
            source="pdf",
            error=(
                "Could not extract text from this PDF. "
                "The file may be image-based (scanned). "
                "Try uploading a Word document or JSON file instead."
            ),
        )


def _extract_docx_text(doc) -> str:
    """Extract text from DOCX preserving structure as markdown-like format.

    Detects headings, list items (via style name, XML numPr, or left indent),
    bold-only paragraphs (treated as section headers), and tables. The output
    is a markdown-ish string that gives Claude enough structural cues to
    correctly identify bullets, section boundaries, and entry titles.
    """
    WML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # ── Heading styles ────────────────────────────────────────────
        if "Heading" in style_name:
            level = 1
            for char in style_name:
                if char.isdigit():
                    level = int(char)
                    break
            parts.append(f"\n{'#' * min(level + 1, 4)} {text}\n")
            continue

        # ── List detection ────────────────────────────────────────────
        is_list = False
        indent_level = 0

        # Method 1: style name contains list-related keywords
        list_keywords = ["list", "bullet", "numbered"]
        if any(kw in style_name.lower() for kw in list_keywords):
            is_list = True

        # Method 2: XML numPr (numbering properties) — catches custom
        # styled lists that don't use a "List …" style name.
        pPr = para._p.pPr
        if pPr is not None:
            numPr = pPr.find(f".//{WML_NS}numPr")
            if numPr is not None:
                is_list = True
                ilvl = numPr.find(f"{WML_NS}ilvl")
                if ilvl is not None:
                    val = ilvl.get(f"{WML_NS}val")
                    if val is not None:
                        indent_level = int(val)

        # Method 3: left indent on non-list-styled but indented paragraphs
        if not is_list and para.paragraph_format.left_indent:
            emu_per_indent = 457200  # 0.5 inch in EMUs
            indent_level = max(
                0, int(para.paragraph_format.left_indent / emu_per_indent)
            )
            if indent_level > 0:
                is_list = True

        if is_list:
            prefix = "  " * indent_level + "- "
            parts.append(f"{prefix}{text}")
            continue

        # ── Bold-only paragraphs → likely section header / entry title ─
        runs = para.runs
        if runs and all(r.bold for r in runs if r.text.strip()):
            parts.append(f"**{text}**")
            continue

        # ── Regular paragraph ─────────────────────────────────────────
        parts.append(text)

    # ── Tables ────────────────────────────────────────────────────────
    for table in doc.tables:
        parts.append("")  # blank line before table
        for row in table.rows:
            cells = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if cells:
                parts.append(cells)

    return "\n".join(parts)


async def extract_from_docx(file_bytes: bytes) -> CVImportResult:
    """Extract text from DOCX via python-docx, then send to Claude for structuring."""
    try:
        doc = Document(BytesIO(file_bytes))
        full_text = _extract_docx_text(doc)
        if not full_text.strip():
            return CVImportResult(
                success=False,
                source="docx",
                error="The DOCX file appears to be empty or contains no extractable text.",
            )

        user_prompt = (
            f"{EXTRACTION_USER_PROMPT}\n\nDocument content:\n---\n{full_text}\n---"
        )

        response = bedrock_client.chat(
            messages=[{"role": "user", "content": user_prompt}],
            system_prompt=EXTRACTION_SYSTEM_PROMPT,
            stream=False,
            max_tokens=8192,
        )
        return _parse_extraction_response(response, source="docx")
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        return CVImportResult(
            success=False,
            source="docx",
            error=f"Failed to extract CV from DOCX: {e}",
        )


async def extract_from_json(file_bytes: bytes) -> CVImportResult:
    """Parse JSON directly, validate against CVFormData schema. No AI needed."""
    try:
        data = json.loads(file_bytes.decode("utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError) as e:
        return CVImportResult(
            success=False,
            source="json",
            error=f"Invalid JSON file: {e}",
        )

    data.pop("templateId", None)

    # Ensure minimum required keys
    data.setdefault("personalInfo", {})
    data.setdefault("workExperience", [])
    data.setdefault("education", [])
    data.setdefault("skills", [])

    try:
        validated = CVFormData(**(data | {"templateId": "_validation"}))
        form_data = validated.model_dump()
        form_data.pop("templateId")
    except ValidationError as e:
        field_errors = [
            f"{'.'.join(str(loc) for loc in err['loc'])}: {err['msg']}"
            for err in e.errors()
        ]
        return CVImportResult(
            success=False,
            source="json",
            error="JSON structure doesn't match the CV schema.",
            warnings=field_errors[:5],
        )

    summary = ImportSummary(
        workEntries=len(form_data.get("workExperience", [])),
        educationEntries=len(form_data.get("education", [])),
        skillCategories=len(form_data.get("skills", [])),
        projects=len(form_data.get("projects") or []),
        awards=len(form_data.get("awards") or []),
        additionalSections=len(form_data.get("additionalSections") or []),
    )

    return CVImportResult(
        success=True,
        form_data=form_data,
        source="json",
        confidence={"overall": "high", "fields": {}},
        summary=summary,
        warnings=None,
    )
